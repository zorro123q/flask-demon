from docx import Document

# 创建一个Word文档
doc = Document()
doc.add_heading('云平台用户体验数据传输优化系统实施效果对比', level=1)

# 添加表格
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '指标'
hdr_cells[1].text = '实施前'
hdr_cells[2].text = '实施后'
hdr_cells[3].text = '改善情况'
hdr_cells[4].text = '地点'

# 添加数据
records = [
    ("视频加载时间", "12秒", "3秒", "减少了75%", "纽约、伦敦、东京"),
    ("播放中断次数(每小时)", "1次", "0.2次", "减少了80%", "纽约、伦敦、东京"),
    ("用户满意度", "78%", "94%", "提升了20.5%", "纽约、伦敦、东京")
]

for rec in records:
    row_cells = table.add_row().cells
    row_cells[0].text = rec[0]
    row_cells[1].text = rec[1]
    row_cells[2].text = rec[2]
    row_cells[3].text = rec[3]
    row_cells[4].text = rec[4]

# 保存文件
file_path = "/mnt/data/优化系统实施效果对比表.docx"
doc.save(file_path)

from docx import Document

# 创建一个Word文档
doc = Document()
doc.add_heading('云平台用户体验数据传输优化系统实施效果对比', level=1)

# 添加表格
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'

# 填充表头
headers = ['指标', '实施前', '实施后', '改善情况']
for col_num, header in enumerate(headers):
    table.cell(0, col_num).text = header

# 填充表格数据
data = [
    ["地点", "纽约、伦敦、东京", "纽约、伦敦、东京", "—"],
    ["视频加载时间", "12秒", "3秒", "减少了75%"],
    ["播放中断次数(每小时)", "1次", "0.2次", "减少了80%"],
    ["用户满意度", "78%", "94%", "提升了20.5%"]
]

for row_num, row_data in enumerate(data, start=1):
    for col_num, cell_data in enumerate(row_data):
        table.cell(row_num, col_num).text = cell_data

# 保存文档
doc_path = "/mnt/data/Cloud_Platform_User_Experience_Data_Transmission_Optimization_System_Performance_Comparison.docx"
doc.save(doc_path)

